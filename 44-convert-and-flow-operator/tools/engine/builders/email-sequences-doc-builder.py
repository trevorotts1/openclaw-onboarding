"""Assemble the v2 email sequences into one human-review document.

Reads all 9 source files from the social-media-tool repo
`docs/email-sequences/` and emits:
  - /tmp/email-sequences-review-doc.docx  (upload artifact -> Google Doc)
  - /tmp/email-sequences-review-doc.md    (plain-text reference / debug)

The .docx carries native heading styles (Title / Heading 1 / Heading 2) so
Google converts it to a properly-outlined Doc. Email bodies are kept
verbatim (the actual copy, not GHL HTML). The `CODE - timing - title`
headings and `Subject A/B:` labels are preserved so the operator's edits can be
parsed back and synced to the source files + GHL drafts.

Usage:
    .venv/bin/python3 builders/email-sequences-doc-builder.py
"""
import re
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _email_sequences_parser import parse_sequence_file

from docx import Document
from docx.shared import Pt

SEQ_DIR = (
    Path.home()
    / "Documents/Tech & Dev/Studio Apps/social-media-tool"
    / "docs/email-sequences"
)
OUT_DOCX = Path("/tmp/email-sequences-review-doc.docx")
OUT_MD = Path("/tmp/email-sequences-review-doc.md")

# Per-section config. `branch` filters WF6's records by code prefix.
WORKFLOWS = [
    dict(n=1, title="WF1 - Course Interest", file="course-interest-rewrite.md",
         branch=None, status="DEPLOYED - GHL draft 30ace83f",
         trigger="Form Submitted = Course Interest", goal="tag: course purchase",
         cadence="10 emails / 14 days"),
    dict(n=2, title="WF5 - High Ticket Interest", file="ht-interest-rewrite.md",
         branch=None, status="DEPLOYED - GHL draft 98fc4aa7",
         trigger="Form Submitted = HT Interest",
         goal="tags: ht-interest-clicked-reserve/-deposit/-consult, purchase:complete",
         cadence="5 emails + 1 SMS / 8 days"),
    dict(n=3, title="WF6-A - Post-Call Sales: HTBO", file="post-call-sales-sequence.md",
         branch="A", status="DEPLOYED - GHL draft ef7309b6",
         trigger="tag: postcall-sequence-htbo", goal="tag: purchase:complete",
         cadence="5 emails / 7 days"),
    dict(n=4, title="WF6-B - Post-Call Sales: LGI", file="post-call-sales-sequence.md",
         branch="B", status="DEPLOYED - GHL draft 4dc0d028",
         trigger="tag: postcall-sequence-lgi", goal="tag: purchase:complete",
         cadence="6 emails / 9 days"),
    dict(n=5, title="WF6-C - Post-Call Sales: AIA", file="post-call-sales-sequence.md",
         branch="C", status="DEPLOYED - GHL draft 67fb1480",
         trigger="tag: postcall-sequence-aia", goal="tag: purchase:complete",
         cadence="6 emails / 9 days"),
    dict(n=6, title="Call Interest", file="call-interest-rewrite.md",
         branch=None, status="DEFERRED - not yet deployed",
         trigger="Form Submitted = Call Interest", goal="tag: call booked",
         cadence="5 emails / 7 days"),
    dict(n=7, title="Generic Call Reminders", file="generic-call-reminders.md",
         branch=None, status="DEFERRED - not yet deployed",
         trigger="tag: call booked", goal="call completed / purchase:complete",
         cadence="3 reminders + 1 no-show, appointment-relative"),
    dict(n=8, title="WF2 - HT Call Reminders", file="ht-schedule-call-patches.md",
         branch=None, status="DEFERRED - not yet deployed",
         trigger="tag: high ticket call booked",
         goal="call completed / purchase:complete",
         cadence="3 reminders + 1 no-show, appointment-relative"),
    dict(n=9, title="WF4 - Call Funnel Sunset", file="call-funnel-patches.md",
         branch=None, status="PLAYBOOK - no email copy to review",
         trigger="n/a", goal="n/a", cadence="n/a"),
]

HOW_TO = [
    "Edit subject lines and body copy directly, inline.",
    "KEEP the structure intact: the email headings (CODE - timing - title) and "
    "the Subject A: / Subject B: labels. Edits sync back to the source files "
    "and the GHL drafts by matching on those markers - if a heading or label "
    "is deleted, that email cannot be matched.",
    "Body copy is shown exactly as it sends. Lines beginning >> are "
    "call-to-action buttons; the URL follows in parentheses. {{contact.first_name}} "
    "and similar are merge tags - leave them as literal text.",
    "To flag something without changing it, use a Google Docs comment.",
    'When you are done, tell your assistant "sync the email doc" and the edits flow '
    "back into the source files and the GHL drafts.",
]

DEPLOY_STATUS = [
    "5 sequences are deployed to GHL as DRAFTs (nothing sends until "
    'published), in the folder "Email Sequences v2" '
    "(542e0d93-2425-439d-aeab-0b83a87cba00):",
    "WF1 Course Interest - draft 30ace83f-a7ec-4683-bed3-07843cb4e0e8",
    "WF5 HT Interest - draft 98fc4aa7-e63b-4cdd-a362-d6aa0836924a",
    "WF6-A Post-Call HTBO - draft ef7309b6-acba-4f97-991a-47f2582c9161",
    "WF6-B Post-Call LGI - draft 4dc0d028-6270-4bae-96ef-e851463cdcba",
    "WF6-C Post-Call AIA - draft 67fb1480-a6a1-4f73-826e-30fd9ceac3d2",
    "4 sequences are drafted but NOT yet deployed (blocked on prerequisites): "
    "Call Interest, Generic Call Reminders, WF2 HT Call Reminders, WF4 sunset.",
]

OPEN_ITEMS = [
    "WF1 goal tag. The goal event is built with tag 'course purchase' (with a "
    "space, matching the legacy WF1). The rewrite spec wrote 'course-purchase' "
    "(hyphen). Confirm which tag your purchase automation actually applies.",
    "WF1 E3 testimonial. E3 contains a Daniel Fitzpatrick testimonial still "
    "marked [TESTIMONIAL - operator to confirm/swap]. Confirm the details or swap "
    "in a different testimonial.",
    "WF6 shipped as 3 separate workflows, one per branch (HTBO/LGI/AIA), each "
    "triggered by a tag - not one workflow branching on offer_interest. The "
    "post-call disposition form must apply postcall-sequence-htbo / -lgi / "
    "-aia matching the rep's selection.",
    "Per-email tracking tags (*-sent-N) were left out of the deployed "
    "workflows to keep them clean. Say so if you want them added back.",
    "Footer. Each email carries an address + unsubscribe footer. On test "
    "sends, confirm GHL is not also appending its own (double footer).",
    "Deferred 4 still need: the Call Interest form, calendar-tag verification, "
    "and appointment-time-relative waits before they deploy.",
]

WF4_NOTE = (
    "This is a sunset / migration playbook - it carries no email copy to "
    "review. WF4's audience now flows into the Call Interest sequence "
    "(section 6). See docs/email-sequences/call-funnel-patches.md for the "
    "7-step migration plan."
)

_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


# ── docx rich-text rendering ──────────────────────────────────────────────

def _add_runs(paragraph, text: str):
    """Add text to a paragraph, honoring **bold** spans. Markdown links
    [label](url) are flattened to 'label (url)' first so the URL stays
    visible for review."""
    text = _LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:               # odd chunks were between ** markers
            run.bold = True


def _add_body(doc, body_md: str):
    for raw in body_md.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if stripped.startswith(">"):                       # blockquote / note
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(stripped.lstrip("> ").rstrip())
            run.italic = True
            continue
        p = doc.add_paragraph()
        _add_runs(p, stripped)


def timing_label(rec) -> str:
    m = re.fullmatch(r"D(\d+)", rec.timing)
    return f"Day {m.group(1)}" if m else rec.timing


# ── build ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    doc.add_heading("Email Sequences v2 - Review & Edit", level=0)
    intro = doc.add_paragraph()
    intro.add_run(
        f"Generated {date.today().isoformat()} from docs/email-sequences/. "
        "All 9 v2 sequences in one place for final copy review before the "
        "GHL drafts go live."
    ).italic = True

    doc.add_heading("How to use this doc", level=1)
    for item in HOW_TO:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("GHL deployment status", level=1)
    for line in DEPLOY_STATUS:
        doc.add_paragraph(line)

    doc.add_heading("Open items needing your decision", level=1)
    for i, item in enumerate(OPEN_ITEMS, 1):
        doc.add_paragraph(f"{i}. {item}")

    counts = []
    for wf in WORKFLOWS:
        doc.add_page_break()
        doc.add_heading(f"{wf['n']}. {wf['title']}", level=1)

        meta = doc.add_paragraph()
        meta.add_run("Status: ").bold = True
        meta.add_run(wf["status"])
        if wf["trigger"] != "n/a":
            for label, key in (("Trigger", "trigger"), ("Goal / exit", "goal"),
                               ("Cadence", "cadence")):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(wf[key])

        records = parse_sequence_file(SEQ_DIR / wf["file"])
        if wf["branch"]:
            records = [r for r in records if r.code.startswith(wf["branch"])]

        n_e = sum(1 for r in records if r.kind == "email")
        n_s = sum(1 for r in records if r.kind == "sms")
        counts.append((wf, n_e, n_s))

        if not records:
            doc.add_paragraph(WF4_NOTE)
            continue

        for rec in records:
            label = timing_label(rec)
            head = (f"{rec.code} - {label} - {rec.title}" if label
                    else f"{rec.code} - {rec.title}")
            doc.add_heading(head, level=2)
            if rec.kind == "sms":
                sp = doc.add_paragraph()
                sp.add_run("Type: ").bold = True
                sp.add_run("SMS")
            for lab, val in (("Subject A", rec.subject_a),
                             ("Subject B", rec.subject_b)):
                if val:
                    p = doc.add_paragraph()
                    p.add_run(f"{lab}: ").bold = True
                    p.add_run(val)
            doc.add_paragraph()
            _add_body(doc, rec.body_md)

    doc.save(OUT_DOCX)
    return counts


def build_md(counts):
    """Plain-text reference copy (sync-back debugging)."""
    lines = ["# Email Sequences v2 - Review & Edit", ""]
    for wf, _, _ in counts:
        records = parse_sequence_file(SEQ_DIR / wf["file"])
        if wf["branch"]:
            records = [r for r in records if r.code.startswith(wf["branch"])]
        lines.append(f"# {wf['n']}. {wf['title']}")
        lines.append("")
        for rec in records:
            label = timing_label(rec)
            head = (f"{rec.code} - {label} - {rec.title}" if label
                    else f"{rec.code} - {rec.title}")
            lines.append(f"### {head}")
            if rec.subject_a:
                lines.append(f"**Subject A:** {rec.subject_a}")
            if rec.subject_b:
                lines.append(f"**Subject B:** {rec.subject_b}")
            lines.append("")
            lines.append(rec.body_md)
            lines.append("")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    counts = build()
    build_md(counts)
    total_e = sum(e for _, e, _ in counts)
    total_s = sum(s for _, _, s in counts)
    print(f"Wrote {OUT_DOCX}  ({OUT_DOCX.stat().st_size:,} bytes)")
    print(f"Wrote {OUT_MD}  ({OUT_MD.stat().st_size:,} bytes)")
    print(f"Sections: {len(WORKFLOWS)}   Emails: {total_e}   SMS: {total_s}")
    for wf, n_e, n_s in counts:
        extra = f" + {n_s} SMS" if n_s else ""
        print(f"  {wf['n']}. {wf['title']:34} {n_e} emails{extra}")


if __name__ == "__main__":
    main()
